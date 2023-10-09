from docx import Document

document = Document("output.docx")

for style in document.styles:
    print(style)

document.paragraphs[0].style = document.styles['文档主标题']
document.paragraphs[1].style = document.styles['次标题']
document.paragraphs[2].style = document.styles['版本号']

for p in document.paragraphs:
    if p.style.name == 'Heading 1':
        p.style = document.styles['一级标题']
    if p.style.name == 'Heading 2':
        p.style = document.styles['二级标题']
    if p.style.name == 'Heading 3':
        p.style = document.styles['三级标题']

document.save("output.docx")